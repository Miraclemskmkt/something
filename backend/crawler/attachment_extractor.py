"""附件正文提取：PDF / Word / Excel 分流，提取后同样过入口黑白名单。"""
from __future__ import annotations

import io
import logging
import re
from urllib.parse import urlparse

import httpx

from config import settings
from crawler.anti_crawl import browser_headers, create_http_client, site_root
from crawler.parser import compact_spaced_text

logger = logging.getLogger(__name__)

ATTACHMENT_SUFFIXES = (".pdf", ".doc", ".docx", ".xls", ".xlsx", ".ppt", ".pptx")

try:
    from pypdf import PdfReader
except ImportError:  # pragma: no cover
    PdfReader = None  # type: ignore

try:
    import pdfplumber
except ImportError:  # pragma: no cover
    pdfplumber = None  # type: ignore

try:
    import docx
except ImportError:  # pragma: no cover
    docx = None  # type: ignore

try:
    import openpyxl
except ImportError:  # pragma: no cover
    openpyxl = None  # type: ignore


def attachment_suffix(url: str) -> str:
    path = urlparse(url or "").path.lower().split("?")[0]
    for ext in ATTACHMENT_SUFFIXES:
        if path.endswith(ext):
            return ext
    return ""


def is_attachment_url(url: str) -> bool:
    """URL 是否直接指向附件（非 HTML 通知页）。"""
    if not url:
        return False
    if attachment_suffix(url):
        return True
    path = urlparse(url).path.lower()
    return "/upload/files/" in path and any(ext in path for ext in ATTACHMENT_SUFFIXES)


def _extract_pdf_text(data: bytes) -> str:
    parts: list[str] = []
    if pdfplumber is not None:
        try:
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                for page in pdf.pages:
                    tables = page.extract_tables() or []
                    for table in tables:
                        for row in table:
                            cells = [str(c or "").strip() for c in row]
                            if any(cells):
                                parts.append(" | ".join(c for c in cells if c))
                    text = page.extract_text() or ""
                    if text.strip():
                        parts.append(text)
            merged = "\n".join(parts)
            if len(merged.strip()) >= 50:
                return merged
        except Exception as e:
            logger.debug("pdfplumber extract failed: %s", e)

    if PdfReader is None:
        return ""
    try:
        reader = PdfReader(io.BytesIO(data))
        page_texts = [(page.extract_text() or "") for page in reader.pages]
        return "\n".join(t for t in page_texts if t.strip())
    except Exception as e:
        logger.debug("pypdf extract failed: %s", e)
        return ""


def _extract_docx_text(data: bytes) -> str:
    if docx is None:
        return ""
    try:
        document = docx.Document(io.BytesIO(data))
        parts: list[str] = []
        for para in document.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as e:
        logger.debug("docx extract failed: %s", e)
        return ""


def _extract_xlsx_text(data: bytes, *, max_rows: int = 200) -> str:
    if openpyxl is None:
        return ""
    try:
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts: list[str] = []
        for sheet in wb.worksheets:
            for i, row in enumerate(sheet.iter_rows(values_only=True)):
                if i >= max_rows:
                    break
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    parts.append(" | ".join(cells))
        wb.close()
        return "\n".join(parts)
    except Exception as e:
        logger.debug("xlsx extract failed: %s", e)
        return ""


def extract_text_from_bytes(data: bytes, url: str = "") -> str:
    """按 URL 后缀从二进制内容提取纯文本。"""
    if not data:
        return ""
    ext = attachment_suffix(url)
    raw = ""
    if ext == ".pdf" or data[:5].startswith(b"%PDF-"):
        raw = _extract_pdf_text(data)
    elif ext in (".docx", ".doc"):
        raw = _extract_docx_text(data) if ext == ".docx" else ""
        if not raw and ext == ".doc":
            logger.debug("legacy .doc not supported without antiword: %s", url[:60])
    elif ext in (".xlsx", ".xls"):
        raw = _extract_xlsx_text(data) if ext == ".xlsx" else ""
    if not raw or len(raw.strip()) < 30:
        return ""
    return compact_spaced_text(raw)


def looks_like_html(content: str | bytes) -> bool:
    sample = content[:4000] if isinstance(content, str) else content[:4000].decode("latin-1", errors="ignore")
    return bool(re.search(r"<[a-zA-Z!/?]", sample))


async def fetch_attachment_text(
    url: str,
    *,
    client: httpx.AsyncClient | None = None,
    title: str = "",
) -> str | None:
    """下载附件并提取文本；提取后过 title_filter（正文作为 summary）。"""
    if not url or not url.startswith("http"):
        return None
    if not is_attachment_url(url):
        return None

    async def _do(c: httpx.AsyncClient) -> str | None:
        try:
            resp = await c.get(
                url,
                headers=browser_headers(url, referer=site_root(url)),
                follow_redirects=True,
            )
            if resp.status_code != 200:
                return None
            data = resp.content
            if not data or len(data) > settings.pdf_max_bytes:
                return None
            text = extract_text_from_bytes(data, url)
            if not text or len(text.strip()) < 50:
                logger.debug("attachment text too short: %s", url[:70])
                return None
            from crawler.noise_filter import passes_title_filter

            probe_title = title or text[:120]
            if not passes_title_filter(probe_title, text[:800]):
                logger.info("附件正文未过入口过滤: %s", url[:70])
                return None
            logger.info("附件正文提取成功 (%d 字): %s", len(text), url[:70])
            return text
        except Exception as e:
            logger.debug("attachment fetch error %s: %s", url[:70], e)
            return None

    if client is not None:
        return await _do(client)
    async with create_http_client() as c:
        return await _do(c)
